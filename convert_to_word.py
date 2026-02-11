#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert PANDUAN_LENGKAP.md to Word Document"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Baca file markdown
with open('PANDUAN_LENGKAP.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Buat dokumen Word
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Function untuk add heading dengan styling
def add_heading(text, level):
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        heading.runs[0].font.size = Pt(20)
        heading.runs[0].bold = True
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        heading.runs[0].font.size = Pt(16)
    return heading

# Parse markdown dan convert ke Word
lines = content.split('\n')
i = 0
in_code_block = False
code_lines = []
in_table = False
table_data = []

while i < len(lines):
    line = lines[i]
    
    # Handle code blocks
    if line.startswith('```'):
        if in_code_block:
            # End code block
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'Intense Quote'
            p.paragraph_format.left_indent = Inches(0.5)
            code_lines = []
            in_code_block = False
        else:
            # Start code block
            in_code_block = True
            code_lines = []
        i += 1
        continue
    
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue
    
    # Handle headings
    if line.startswith('# '):
        add_heading(line[2:], 1)
    elif line.startswith('## '):
        add_heading(line[3:], 2)
    elif line.startswith('### '):
        add_heading(line[4:], 3)
    elif line.startswith('#### '):
        add_heading(line[5:], 4)
    
    # Handle horizontal rules
    elif line.strip() == '---':
        doc.add_paragraph('_' * 50)
    
    # Handle lists
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        # Clean markdown formatting
        text = text.replace('**', '').replace('`', '').replace('✅', '✓').replace('❌', '✗').replace('⚠️', '!')
        p = doc.add_paragraph(text, style='List Bullet')
    
    elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
        text = line.strip()[3:].strip()
        text = text.replace('**', '').replace('`', '').replace('✅', '✓').replace('❌', '✗')
        p = doc.add_paragraph(text, style='List Number')
    
    # Handle tables (simple version)
    elif '|' in line and line.strip().startswith('|'):
        if not in_table:
            in_table = True
            table_data = []
        
        # Parse row
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        table_data.append(cells)
        
        # Check if next line is separator or end of table
        if i + 1 < len(lines):
            next_line = lines[i + 1]
            if not ('|' in next_line and next_line.strip().startswith('|')):
                # End of table - create it
                if len(table_data) > 1 and not all('-' in cell for cell in table_data[1]):
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            if row_idx == 0:
                                # Header row
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                
                in_table = False
                table_data = []
    
    # Handle regular paragraphs
    elif line.strip():
        if not in_table:
            # Clean markdown formatting
            text = line.replace('**', '').replace('`', '').replace('✅', '✓').replace('❌', '✗').replace('⚠️', '!')
            
            # Skip separator lines
            if not (all(c in '-=_' for c in text.strip()) and len(text.strip()) > 10):
                p = doc.add_paragraph(text)
    
    # Handle empty lines
    else:
        if not in_code_block and not in_table:
            doc.add_paragraph()
    
    i += 1

# Save document
output_file = 'PANDUAN_LENGKAP.docx'
doc.save(output_file)
print(f'✅ Dokumen Word berhasil dibuat: {output_file}')
print(f'   Lokasi: {output_file}')
